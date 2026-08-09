from __future__ import annotations

import time
from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient

from app.main import app
from app.services.storage import TEMP_ROOT

client = TestClient(app)


def _make_docx(path: Path) -> None:
    doc = Document()
    doc.add_heading("API Teste", level=1)
    doc.add_paragraph("Conteúdo da API.")
    doc.save(str(path))


def _wait_done(task_id: str, timeout: float = 30.0) -> dict:
    """Polling do status até a conversão terminar."""
    deadline = time.monotonic() + timeout
    while time.monotonic() < deadline:
        response = client.get(f"/api/status/{task_id}")
        assert response.status_code == 200
        data = response.json()
        if data["status"] in ("done", "error"):
            return data
        time.sleep(0.2)
    raise AssertionError("Conversão não terminou a tempo")


def test_health() -> None:
    response = client.get("/api/health")
    assert response.status_code == 200
    assert response.json() == {"status": "ok"}


def test_index_page() -> None:
    response = client.get("/")
    assert response.status_code == 200
    assert "Conversor de Arquivos para Markdown" in response.text


def test_robots_txt() -> None:
    response = client.get("/robots.txt")
    assert response.status_code == 200
    assert "User-agent: *" in response.text
    assert "/sitemap.xml" in response.text


def test_sitemap_xml() -> None:
    response = client.get("/sitemap.xml")
    assert response.status_code == 200
    assert "urlset" in response.text
    assert "<loc>" in response.text
    assert "/blog" in response.text
    assert "<lastmod>" in response.text


def test_blog_list() -> None:
    response = client.get("/blog")
    assert response.status_code == 200
    assert "Blog e" in response.text
    assert "post-card" in response.text


def test_blog_article() -> None:
    response = client.get("/blog/como-converter-word-para-markdown")
    assert response.status_code == 200
    assert "Como converter Word para Markdown" in response.text
    assert "BreadcrumbList" in response.text
    assert "post-content" in response.text


def test_blog_article_not_found() -> None:
    response = client.get("/blog/artigo-inexistente")
    assert response.status_code == 404


def test_privacy_page() -> None:
    response = client.get("/privacy")
    assert response.status_code == 200
    assert "Política de Privacidade" in response.text
    assert "cookies" in response.text.lower()


def test_terms_page() -> None:
    response = client.get("/terms")
    assert response.status_code == 200
    assert "Termos de Uso" in response.text


def test_contact_page() -> None:
    response = client.get("/contact")
    assert response.status_code == 200
    assert "Contato" in response.text


def test_ads_txt_default() -> None:
    response = client.get("/ads.txt")
    assert response.status_code == 200
    assert "AdSense não configurado" in response.text


def test_cookie_banner_present() -> None:
    response = client.get("/")
    assert response.status_code == 200
    assert "cookie-banner" in response.text
    assert "cookie-accept" in response.text


def test_static_cache_headers() -> None:
    response = client.get("/static/css/style.css")
    assert response.status_code == 200
    assert "max-age=3600" in response.headers.get("Cache-Control", "")


def test_dynamic_pages_no_cache() -> None:
    for path in ("/", "/blog"):
        response = client.get(path)
        assert response.status_code == 200
        assert response.headers.get("Cache-Control") == "no-cache"


def test_sitemap_xml_no_cache() -> None:
    response = client.get("/sitemap.xml")
    assert response.status_code == 200
    assert response.headers.get("Cache-Control") == "no-cache"


def test_convert_and_download(tmp_path: Path) -> None:
    source = tmp_path / "api.docx"
    _make_docx(source)

    with source.open("rb") as f:
        response = client.post(
            "/api/convert",
            files={"file": ("api.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )

    assert response.status_code == 200
    data = response.json()
    assert data["status"] == "processing"
    assert data["filename"].endswith(".md")
    assert data["status_url"] == f"/api/status/{data['task_id']}"
    assert data["download_url"] == f"/api/download/{data['task_id']}"

    job_dir = TEMP_ROOT / data["task_id"]
    assert job_dir.is_dir()

    status = _wait_done(data["task_id"])
    assert status["status"] == "done"
    assert status["filename"].endswith(".md")
    assert "# API Teste" in status["markdown"]
    assert status["download_url"] == f"/api/download/{data['task_id']}"

    download = client.get(f"/api/download/{data['task_id']}")
    assert download.status_code == 200
    assert b"API Teste" in download.content

    assert not job_dir.is_dir(), "Diretório temporário deve ser excluído após o download"


def test_convert_unknown_task() -> None:
    response = client.get("/api/download/inexistente")
    assert response.status_code == 404


def test_status_unknown_task() -> None:
    response = client.get("/api/status/inexistente")
    assert response.status_code == 404


def test_reject_bad_extension() -> None:
    response = client.post(
        "/api/convert",
        files={"file": ("malware.exe", b"x", "application/octet-stream")},
    )
    assert response.status_code == 400


def test_reject_empty_file() -> None:
    response = client.post(
        "/api/convert",
        files={"file": ("vazio.docx", b"", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert response.status_code == 400
