from __future__ import annotations

from pathlib import Path

import pytest

from app.converters.docx import DocxConverter
from app.converters.xlsx import XlsxConverter
from app.services.libreoffice import (
    LEGACY_TARGETS,
    LibreOfficeUnavailableError,
    _find_soffice,
)
from app.services.magic import sniff_format


def _has_libreoffice() -> bool:
    try:
        _find_soffice()
        return True
    except LibreOfficeUnavailableError:
        return False


def _make_docx(path: Path) -> None:
    from docx import Document

    doc = Document()
    doc.add_heading("Legado", level=1)
    doc.add_paragraph("Conteúdo convertido via LibreOffice.")
    doc.save(str(path))


def _make_xlsx(path: Path) -> None:
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Dados"
    ws.append(["Nome", "Valor"])
    ws.append(["A", 1])
    ws.append(["B", 2])
    wb.save(str(path))


def _soffice() -> str:
    return _find_soffice()


def _to_legacy(source: Path, out: Path, target: str) -> Path:
    import subprocess

    out.mkdir(parents=True, exist_ok=True)
    subprocess.run(
        [_soffice(), "--headless", "--convert-to", target, "--outdir", str(out), str(source)],
        check=True,
        capture_output=True,
    )
    return out / f"{source.stem}.{target}"


@pytest.mark.skipif(not _has_libreoffice(), reason="LibreOffice não instalado")
def test_doc_to_docx_conversion(tmp_path: Path) -> None:
    src = tmp_path / "input.docx"
    _make_docx(src)
    legacy = _to_legacy(src, tmp_path / "legacy", "doc")

    assert sniff_format(legacy) == "doc"

    converted = tmp_path / "modern"
    import subprocess

    result = subprocess.run(
        [
            _soffice(),
            "--headless",
            "--convert-to",
            "docx",
            "--outdir",
            str(converted),
            str(legacy),
        ],
        check=True,
        capture_output=True,
    )
    assert result.returncode == 0
    docx_path = converted / "input.docx"
    assert docx_path.exists()
    assert sniff_format(docx_path) == "docx"

    md = DocxConverter().convert(docx_path, converted / "out").markdown
    assert "Legado" in md


@pytest.mark.skipif(not _has_libreoffice(), reason="LibreOffice não instalado")
def test_xls_to_xlsx_conversion(tmp_path: Path) -> None:
    src = tmp_path / "dados.xlsx"
    _make_xlsx(src)
    legacy = _to_legacy(src, tmp_path / "legacy", "xls")

    assert sniff_format(legacy) == "xls"

    md = XlsxConverter().convert(legacy, tmp_path / "out").markdown
    assert "Nome" in md


@pytest.mark.skipif(not _has_libreoffice(), reason="LibreOffice não instalado")
def test_legacy_targets_mapping() -> None:
    assert LEGACY_TARGETS == {
        "doc": "docx",
        "xls": "xlsx",
        "ppt": "pptx",
        "odt": "docx",
        "ods": "xlsx",
        "odp": "pptx",
    }


def test_soffice_detected_or_raises() -> None:
    if _has_libreoffice():
        assert Path(_find_soffice()).is_file()
    else:
        with pytest.raises(LibreOfficeUnavailableError):
            _find_soffice()
