from __future__ import annotations

from pathlib import Path

from docx import Document
from fastapi import FastAPI
from fastapi.testclient import TestClient

from app.main import app
from app.middleware import RateLimitMiddleware
from app.services.magic import sniff_format

client = TestClient(app)

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _make_docx(path: Path) -> None:
    doc = Document()
    doc.add_heading("Segurança", level=1)
    doc.add_paragraph("Conteúdo válido.")
    doc.save(str(path))


def test_reject_mismatched_content(tmp_path: Path) -> None:
    fake = tmp_path / "fake.docx"
    fake.write_text("isto não é um docx de verdade", encoding="utf-8")

    with fake.open("rb") as f:
        response = client.post("/api/convert", files={"file": ("fake.docx", f, DOCX_MIME)})

    assert response.status_code == 400
    assert "conteúdo" in response.json()["detail"].lower()


def test_reject_renamed_extension(tmp_path: Path) -> None:
    source = tmp_path / "real.docx"
    _make_docx(source)

    with source.open("rb") as f:
        response = client.post(
            "/api/convert",
            files={"file": ("malware.exe", f, "application/octet-stream")},
        )

    assert response.status_code == 400


def test_sniff_detects_docx(tmp_path: Path) -> None:
    source = tmp_path / "doc.docx"
    _make_docx(source)
    assert sniff_format(source) == "docx"


def test_sniff_rejects_plain_text(tmp_path: Path) -> None:
    source = tmp_path / "doc.txt"
    source.write_text("apenas texto", encoding="utf-8")
    assert sniff_format(source) is None


def test_rate_limit() -> None:
    test_app = FastAPI()
    test_app.add_middleware(RateLimitMiddleware, max_requests=3, window_seconds=60)

    @test_app.post("/api/convert")
    async def fake_convert():
        return {"ok": True}

    test_client = TestClient(test_app)
    for _ in range(3):
        response = test_client.post("/api/convert")
        assert response.status_code == 200

    blocked = test_client.post("/api/convert")
    assert blocked.status_code == 429
