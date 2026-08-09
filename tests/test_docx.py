from __future__ import annotations

from pathlib import Path

from docx import Document

from app.converters.docx import DocxConverter


def _make_docx(path: Path) -> None:
    doc = Document()
    doc.add_heading("Título 1", level=1)

    paragraph = doc.add_paragraph()
    paragraph.add_run("Texto em ")
    bold = paragraph.add_run("negrito")
    bold.bold = True
    paragraph.add_run(".")

    doc.add_paragraph("Parágrafo simples.")
    doc.add_paragraph("Item 1", style="List Bullet")
    doc.add_paragraph("Item 2", style="List Bullet")

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Coluna A"
    table.cell(0, 1).text = "Coluna B"
    table.cell(1, 0).text = "Valor 1"
    table.cell(1, 1).text = "Valor 2"

    doc.save(str(path))


def test_convert_docx(tmp_path: Path) -> None:
    source = tmp_path / "documento.docx"
    _make_docx(source)

    result = DocxConverter().convert(source, tmp_path / "output")

    assert result.markdown
    assert "# Título 1" in result.markdown
    assert "negrito" in result.markdown
    assert "Parágrafo simples" in result.markdown
    assert "Item 1" in result.markdown
    assert "Coluna A" in result.markdown
    assert "Valor 1" in result.markdown
